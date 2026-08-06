# -*- coding: utf-8 -*-
"""
线口监测 · 云端版(GitHub Actions 定时运行)
流程:抓取RSS/HN/Bing新闻 → 通义千问总结分级 → 生成JSON(网页用)+Word简报 → 更新manifest
环境变量:
  DASHSCOPE_API_KEY  百炼API Key(必填,否则降级为无摘要模式)
  QWEN_MODEL         可选,默认 qwen-plus
"""
import os, re, json, time, datetime, zoneinfo, urllib.parse
import requests, feedparser

BASE = os.path.dirname(os.path.abspath(__file__))
WEB = os.path.join(BASE, "web")
DATA = os.path.join(WEB, "data")
TZ = zoneinfo.ZoneInfo("Asia/Shanghai")
NOW = datetime.datetime.now(TZ)
UA = {"User-Agent": "Mozilla/5.0 (Macintosh; Intel Mac OS X 10_15_7) AppleWebKit/537.36 (KHTML, like Gecko) Chrome/120.0 Safari/537.36"}

BEATS = ["AI/大模型", "机器人/AI硬件", "无人驾驶/低空/智慧出行", "电商/直播/即时零售", "游戏/电竞/平台公司", "消费/以旧换新/新餐饮"]

# RSS 信源(尽力而为,失败自动跳过)
FEEDS = [
    ("36氪", "https://36kr.com/feed"),
    ("IT之家", "https://www.ithome.com/rss/"),
    ("GameLook", "http://www.gamelook.com.cn/feed"),
]

# Bing 新闻检索词(每线口兜底)
QUERIES = {
    "AI/大模型": ["大模型 发布", "DeepSeek OR 通义 OR 智谱"],
    "机器人/AI硬件": ["人形机器人 OR 具身智能", "AI眼镜"],
    "无人驾驶/低空/智慧出行": ["无人驾驶 OR Robotaxi", "低空经济 OR eVTOL"],
    "电商/直播/即时零售": ["外卖 OR 即时零售", "跨境电商 OR Temu OR SHEIN"],
    "游戏/电竞/平台公司": ["游戏出海 OR 版号", "电竞 OR 小游戏"],
    "消费/以旧换新/新餐饮": ["以旧换新 OR 国补", "新茶饮 OR 餐饮 消费"],
}

# 重点实体(合并检索,减少请求数)
ENTITY_QUERIES = [
    "DeepSeek OR 通义千问 OR 智谱 OR Kimi 最新",
    "宇树 OR 优必选 OR 智元机器人 OR 傅利叶 最新",
    "文远知行 OR 小马智行 OR 萝卜快跑 OR 亿航 最新",
    "京东外卖 OR 淘宝闪购 OR 美团闪购 OR 前置仓 最新",
    "SHEIN OR Temu OR TikTok Shop OR 广交会 最新",
    "腾讯 OR 网易 OR 米哈游 OR 微信小游戏 最新",
    "瑞幸 OR 新茶饮 OR 唯品会 OR 以旧换新 最新",
]

LOCAL_WORDS = ["广州", "深圳", "广东", "大湾区", "南沙", "黄埔", "天河"]
HN_TECH_WORDS = ["AI", "LLM", "GPT", "model", "robot", "autonomous", "drone", "EV", "chip", "agent", "open source"]
WINDOW_H = 48  # 抓取窗口(小时)

stats = {"ok": [], "fail": []}


def log(*a):
    print("[monitor]", *a, flush=True)


def within_window(ts):
    if not ts:
        return False
    dt = datetime.datetime.fromtimestamp(ts, TZ)
    return (NOW - dt).total_seconds() <= WINDOW_H * 3600


def fmt_time(ts):
    if not ts:
        return ""
    return datetime.datetime.fromtimestamp(ts, TZ).strftime("%m-%d %H:%M")


# ---------- 抓取 ----------
def fetch_rss():
    items = []
    for name, url in FEEDS:
        try:
            r = requests.get(url, headers=UA, timeout=20)
            if r.status_code != 200:
                raise RuntimeError("HTTP %s" % r.status_code)
            fp = feedparser.parse(r.content)
            if not fp.entries:
                raise RuntimeError("no entries")
            n = 0
            for e in fp.entries:
                ts = time.mktime(e.published_parsed) if getattr(e, "published_parsed", None) else (
                    time.mktime(e.updated_parsed) if getattr(e, "updated_parsed", None) else 0)
                if not within_window(ts):
                    continue
                items.append({"title": re.sub(r"\s+", " ", e.get("title", "")).strip(),
                              "url": e.get("link", ""), "source": name, "ts": ts,
                              "snippet": re.sub(r"<[^>]+>", "", e.get("summary", ""))[:200]})
                n += 1
            stats["ok"].append("%s(%d)" % (name, n))
        except Exception as ex:
            stats["fail"].append("%s(%s)" % (name, str(ex)[:40]))
            log("RSS失败:", name, ex)
    return items


def fetch_hn():
    from concurrent.futures import ThreadPoolExecutor
    items = []
    try:
        ids = requests.get("https://hacker-news.firebaseio.com/v0/topstories.json", timeout=15).json()[:30]

        def one(i):
            try:
                return requests.get("https://hacker-news.firebaseio.com/v0/item/%s.json" % i, timeout=8).json()
            except Exception:
                return None

        with ThreadPoolExecutor(max_workers=12) as ex:
            ds = list(ex.map(one, ids))
        for d in ds:
            if not d:
                continue
            title = d.get("title", "")
            if d.get("score", 0) < 80 or not any(w.lower() in title.lower() for w in HN_TECH_WORDS):
                continue
            if not within_window(d.get("time")):
                continue
            items.append({"title": title, "url": d.get("url") or "https://news.ycombinator.com/item?id=%s" % d.get("id"),
                          "source": "Hacker News", "ts": d.get("time"), "snippet": "score %s" % d.get("score")})
        stats["ok"].append("HackerNews(%d)" % len(items))
    except Exception as ex:
        stats["fail"].append("HackerNews(%s)" % str(ex)[:40])
        log("HN失败:", ex)
    return items


def fetch_bing(query):
    url = "https://www.bing.com/news/search?q=%s&format=rss" % urllib.parse.quote(query)
    r = requests.get(url, headers=UA, timeout=10)
    fp = feedparser.parse(r.content)
    out = []
    for e in fp.entries:
        ts = time.mktime(e.published_parsed) if getattr(e, "published_parsed", None) else 0
        if not within_window(ts):
            continue
        src = e.get("source", {}).get("title", "") if isinstance(e.get("source"), dict) else ""
        out.append({"title": re.sub(r"\s+", " ", e.get("title", "")).strip(), "url": e.get("link", ""),
                    "source": src or "Bing新闻", "ts": ts,
                    "snippet": re.sub(r"<[^>]+>", "", e.get("summary", ""))[:200]})
    return out


def fetch_bing_all():
    items = []
    deadline = time.time() + 240  # 检索阶段总预算4分钟,超时就跳过剩余,保证整体不超时
    for beat, qs in QUERIES.items():
        for q in qs:
            if time.time() > deadline:
                log("Bing检索预算用完,跳过剩余")
                break
            try:
                got = fetch_bing(q)
                for g in got:
                    g["hint_beat"] = beat
                items += got
                time.sleep(0.3)
            except Exception as ex:
                log("Bing失败:", q, ex)
    for q in ENTITY_QUERIES:
        if time.time() > deadline:
            log("Bing检索预算用完,跳过剩余实体")
            break
        try:
            items += fetch_bing(q)
            time.sleep(0.3)
        except Exception as ex:
            log("Bing失败:", q, ex)
    if items:
        stats["ok"].append("Bing新闻(%d)" % len(items))
    else:
        stats["fail"].append("Bing新闻(无结果)")
    return items


def dedupe(items):
    seen, out = set(), []
    for it in items:
        if not it.get("title"):
            continue
        key = re.sub(r"[\s\W]", "", it["title"])[:24]
        if key in seen:
            continue
        seen.add(key)
        out.append(it)
    return out


# ---------- 通义千问 ----------
def llm_bucket(pool):
    """调用通义,把候选池分入六大线口并分级总结;失败返回None走兜底。"""
    key = os.environ.get("DASHSCOPE_API_KEY", "").strip()
    if not key:
        log("未配置DASHSCOPE_API_KEY,使用无摘要兜底模式")
        return None
    from openai import OpenAI
    client = OpenAI(api_key=key, base_url="https://dashscope.aliyuncs.com/compatible-mode/v1",
                    timeout=120, max_retries=1)
    model = os.environ.get("QWEN_MODEL", "qwen-plus")
    cand = [{"t": it["title"], "s": it["source"], "tm": fmt_time(it["ts"]), "u": it["url"],
             "d": it.get("snippet", "")[:120]} for it in pool[:80]]
    prompt = (
        "你是羊城晚报记者的新闻线索筛选助手。以下是候选新闻(JSON数组,t=标题,s=来源,tm=时间,u=链接,d=摘要片段)。\n"
        "请完成:1)挑选最有新闻价值的线索,总数不超过30条;2)每条归入线口:%s;"
        "3)分级:A=突发重大值得立即跟进,B=选题储备,C=背景参考;4)写2-3句中文摘要;5)涉及广州/广东/大湾区的local=true;6)A/B级给一句角度建议。\n"
        "严格只输出JSON,格式:{\"items\":[{\"beat\":线口,\"title\":原标题,\"url\":链接,\"source\":来源,\"time\":时间,\"summary\":摘要,\"grade\":\"A|B|C\",\"local\":bool,\"angle\":建议或空字符串}]}\n"
        "候选:%s" % ("、".join(BEATS), json.dumps(cand, ensure_ascii=False))
    )
    try:
        resp = client.chat.completions.create(
            model=model, messages=[{"role": "user", "content": prompt}],
            response_format={"type": "json_object"}, temperature=0.2)
        text = resp.choices[0].message.content
        data = json.loads(text)
        items = [it for it in data.get("items", []) if it.get("beat") in BEATS and it.get("title")]
        log("通义返回 %d 条分级线索" % len(items))
        return items or None
    except Exception as ex:
        log("通义调用失败:", ex)
        return None


def fallback_bucket(pool):
    """无API Key时的兜底:按关键词粗分线口,不写摘要。"""
    rules = {
        "AI/大模型": ["大模型", "智能体", "AI", "DeepSeek", "通义", "智谱", "Kimi", "OpenAI", "算力"],
        "机器人/AI硬件": ["机器人", "具身", "AI眼镜", "宇树", "优必选", "robot"],
        "无人驾驶/低空/智慧出行": ["无人驾驶", "Robotaxi", "自动驾驶", "低空", "eVTOL", "无人机", "网约车", "智驾"],
        "电商/直播/即时零售": ["电商", "直播", "外卖", "即时零售", "跨境", "Temu", "SHEIN", "广交会", "闪购"],
        "游戏/电竞/平台公司": ["游戏", "电竞", "版号", "小游戏", "腾讯", "网易", "米哈游"],
        "消费/以旧换新/新餐饮": ["以旧换新", "国补", "消费", "茶饮", "餐饮", "免税", "零售"],
    }
    out, seen = [], set()
    for it in pool:
        beat = next((b for b, ws in rules.items() if any(w.lower() in it["title"].lower() for w in ws)), None)
        if not beat:
            continue
        k = it["title"][:24]
        if k in seen:
            continue
        seen.add(k)
        out.append({"beat": beat, "title": it["title"], "url": it["url"], "source": it["source"],
                    "time": fmt_time(it["ts"]), "summary": "", "grade": "C",
                    "local": any(w in it["title"] for w in LOCAL_WORDS), "angle": ""})
        if len(out) >= 30:
            break
    return out


# ---------- Word ----------
def build_docx(sections, flash, path, title):
    from docx import Document
    from docx.shared import Pt, RGBColor
    from docx.oxml.ns import qn
    from docx.oxml import OxmlElement
    BRAND, GRAY = RGBColor(0x41, 0xD8, 0x7E), RGBColor(0x88, 0x88, 0x88)

    def cn(x, name="宋体", size=10.5):
        x.font.name = "Times New Roman"
        x.font.size = Pt(size)
        x.element.get_or_add_rPr().get_or_add_rFonts().set(qn("w:eastAsia"), name)
        return x

    def link(par, url, text):
        rid = par.part.relate_to(url, "http://schemas.openxmlformats.org/officeDocument/2006/relationships/hyperlink", is_external=True)
        a = OxmlElement("w:hyperlink"); a.set(qn("r:id"), rid)
        r = OxmlElement("w:r"); rpr = OxmlElement("w:rPr")
        c = OxmlElement("w:color"); c.set(qn("w:val"), "0563C1"); rpr.append(c)
        u = OxmlElement("w:u"); u.set(qn("w:val"), "single"); rpr.append(u)
        r.append(rpr)
        t = OxmlElement("w:t"); t.text = text; r.append(t)
        a.append(r); par._p.append(a)

    doc = Document()
    cn(doc.styles["Normal"])
    h = doc.add_heading(level=1); cn(h.add_run(title), "黑体", 16).font.color.rgb = BRAND
    sub = doc.add_paragraph()
    cn(sub.add_run("监测窗口:近%d小时 | 生成时间:%s" % (WINDOW_H, NOW.strftime("%Y-%m-%d %H:%M"))), size=9).font.color.rgb = GRAY

    h = doc.add_heading(level=2); cn(h.add_run("今日要闻速览"), "黑体", 14).font.color.rgb = BRAND
    for i, f in enumerate(flash[:8], 1):
        p = doc.add_paragraph()
        cn(p.add_run("%d. " % i)).bold = True
        if f.get("local"):
            rl = p.add_run("【本地】"); cn(rl).bold = True; rl.font.color.rgb = BRAND
        cn(p.add_run("%s——%s(%s)" % (f["title"], f["summary"] or "详见下文", f["source"])))

    for beat in BEATS:
        h = doc.add_heading(level=2); cn(h.add_run(beat), "黑体", 14)
        its = sections.get(beat, [])
        if not its:
            doc.add_paragraph("本线口暂无重要更新")
            continue
        for it in its:
            h3 = doc.add_heading(level=3)
            tag = {"A": "【A级·线索", "B": "【B级·关注", "C": "【C级·参考"}.get(it["grade"], "【C级·参考")
            tag += "|本地】" if it.get("local") else "】"
            r1, r2 = h3.add_run(tag), h3.add_run(it["title"])
            for r in (r1, r2):
                cn(r, "黑体", 12)
                r.bold = it["grade"] in ("A", "B")
                if it["grade"] == "A":
                    r.font.color.rgb = BRAND
                elif it["grade"] == "C":
                    r.bold = False; r.font.color.rgb = GRAY
            p = doc.add_paragraph(); cn(p.add_run("来源:%s | 时间:%s" % (it["source"], it.get("time", ""))), size=9).font.color.rgb = GRAY
            if it.get("summary"):
                doc.add_paragraph(it["summary"])
            if it.get("angle"):
                pa = doc.add_paragraph(); cn(pa.add_run("角度建议:" + it["angle"])).italic = True
            if it.get("url"):
                pl = doc.add_paragraph(); cn(pl.add_run("链接:"), size=9)
                link(pl, it["url"], it["title"])

    h = doc.add_heading(level=2); cn(h.add_run("信源状态"), "黑体", 14)
    p = doc.add_paragraph(); cn(p.add_run("可用:"), size=9).bold = True; cn(p.add_run("、".join(stats["ok"]) or "无"), size=9)
    p = doc.add_paragraph(); cn(p.add_run("暂不可用:"), size=9).bold = True; cn(p.add_run("、".join(stats["fail"]) or "无"), size=9)
    tail = doc.add_paragraph(); cn(tail.add_run("千问办公——专业人士,都用千问办公。"), size=8).font.color.rgb = GRAY
    doc.save(path)


# ---------- 主流程 ----------
def main():
    shift = "早报" if NOW.hour < 12 else "午报"
    slug = NOW.strftime("%Y-%m-%d") + "-" + ("zaobao" if shift == "早报" else "wubao")
    log("班次:", shift, NOW.isoformat())

    pool = dedupe(fetch_rss() + fetch_hn() + fetch_bing_all())
    pool.sort(key=lambda x: -(x.get("ts") or 0))
    log("候选池:", len(pool), "条")
    if not pool:
        log("无任何候选,退出")
        return

    items = llm_bucket(pool) or fallback_bucket(pool)

    sections = {b: [] for b in BEATS}
    order = {"A": 0, "B": 1, "C": 2}
    for it in items:
        if it.get("beat") in sections:
            sections[it["beat"]].append(it)
    for b in sections:
        sections[b].sort(key=lambda x: order.get(x.get("grade"), 3))
    flash = sorted([it for b in BEATS for it in sections[b] if it["grade"] == "A"],
                   key=lambda x: 0 if x.get("local") else 1)

    title = "线口监测%s · %s(%s)" % (shift, NOW.strftime("%Y年%m月%d日"), "周" + "一二三四五六日"[NOW.weekday()])

    os.makedirs(os.path.join(DATA, "briefings"), exist_ok=True)
    os.makedirs(os.path.join(BASE, "briefings"), exist_ok=True)
    docx_name = "线口监测%s-%s.docx" % (shift, NOW.strftime("%Y-%m-%d"))
    docx_rel = "briefings/" + docx_name
    build_docx(sections, flash, os.path.join(BASE, docx_rel), title)

    brief = {"id": slug, "title": title, "shift": shift, "date": NOW.strftime("%Y-%m-%d %H:%M"),
             "docx": "../" + docx_rel, "sections": sections,
             "flash": flash[:8], "stats": stats}
    with open(os.path.join(DATA, "briefings", slug + ".json"), "w", encoding="utf-8") as f:
        json.dump(brief, f, ensure_ascii=False, indent=1)

    mf_path = os.path.join(DATA, "manifest.json")
    manifest = []
    if os.path.exists(mf_path):
        try:
            manifest = json.load(open(mf_path, encoding="utf-8"))
        except Exception:
            manifest = []
    manifest = [m for m in manifest if m["id"] != slug]
    manifest.insert(0, {"id": slug, "title": title, "date": NOW.strftime("%Y-%m-%d %H:%M"),
                        "dataPath": "data/briefings/%s.json" % slug, "docxPath": docx_rel,
                        "top": [f["title"] for f in flash[:4]]})
    json.dump(manifest[:60], open(mf_path, "w", encoding="utf-8"), ensure_ascii=False, indent=1)

    a_cnt = sum(1 for b in BEATS for it in sections[b] if it["grade"] == "A")
    log("完成:A级%d条,总%d条 | Word:%s" % (a_cnt, len(items), docx_rel))


if __name__ == "__main__":
    main()
